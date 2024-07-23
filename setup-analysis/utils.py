from pathlib import Path
from typing import Union, List, Dict, Callable, NamedTuple
import pickle
import matplotlib.pyplot as plt
import seaborn as sns
import inspect
from screeninfo import get_monitors
import shutil
from PIL import Image
import numpy as np
import pandas as pd
import time
import pendulum
from loguru import logger
from functools import wraps
from rich import print as rprint

import screeninfo


def timer(enabled=True):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            if not enabled:
                return func(*args, **kwargs)

            start = time.perf_counter()
            results = func(*args, **kwargs)
            end = time.perf_counter()

            elapsed = pendulum.duration(seconds=end - start)
            elapsed_str = f"{elapsed.total_seconds()} seconds"
            elapsed_str = f"Elapsed time for `{func.__name__}`: {elapsed_str}"

            rprint(f"[green]{elapsed_str}[/green]")
            logger.info(elapsed_str)

            return results

        return wrapper

    return decorator


def get_monitors_info():
    current_monitors = [
        {
            "name": monitor.name if monitor.name else f"Monitor{i+1}",
            "res": (monitor.width, monitor.height),
            "width": monitor.width,
            "height": monitor.height,
            "width_mm": monitor.width_mm,
            "height_mm": monitor.height_mm,
            "primary": monitor.is_primary,
        }
        for i, monitor in enumerate(get_monitors())
    ]

    return current_monitors


def reset_dir(directory: Union[str, Path], keep_stuct: bool = True):
    directory = Path(directory)

    if directory.exists():
        subdirs = [d for d in directory.iterdir() if d.is_dir()]

        for d in subdirs:
            shutil.rmtree(d)

        for subdir in subdirs:
            subdir.mkdir(parents=True, exist_ok=True)
    else:
        directory.mkdir(parents=True, exist_ok=True)


def get_pixel_counts(images: Dict[str, Image.Image]):
    # images = {k: Image.open(v) for k, v in imgs_dict.items()}

    pixel_counts = {}
    for img_name, img in images.items():
        arr = np.asarray(img)
        converted = np.where(arr > 0, 255, 0)
        unique, counts = np.unique(converted, return_counts=True)
        pixel_counts[img_name] = dict(zip(unique, counts))

    df = pd.DataFrame(pixel_counts).T
    df.rename(columns={0: "white", 255: "black"}, inplace=True)
    df.drop(columns=["white"], inplace=True)

    df["diff_ratio"] = round(df["black"] / df["black"].max(), 3)

    df.sort_values("diff_ratio", ascending=False, inplace=True)
    df.reset_index(drop=False, inplace=True)
    df.rename(columns={"index": "img_name"}, inplace=True)

    df["rank"] = df["diff_ratio"].rank(method="dense", ascending=False)

    df.groupby("rank")["black"].mean()

    # converted[:, :, 0][np.where(converted[:, :, 3] == 255)] = 255
    # converted = converted.astype(np.uint8)
    # converted = Image.fromarray(converted)

    return df


def pickle_save(path, data):
    path = Path(path)
    suffix = path.suffix
    if suffix not in [".pickle", ".pkl"]:
        path = path.parent.resolve() / f"{path.stem}.pkl"
    with open(path, "wb") as f:
        pickle.dump(data, f)


def pickle_load(path):
    path = Path(path)
    suffix = path.suffix
    if suffix == "":
        path += ".pickle"

    if not path.exists():
        raise ValueError("File does not exist")

    elif suffix not in [".pickle", ".pkl"]:
        raise ValueError("Not a pickle file")

    if not Path(path).exists():
        raise ValueError("The file does not exist, check the path you provided.")
    with open(path, "rb") as f:
        return pickle.load(f)


def invert_dict(d: Dict) -> Dict:
    return {v: k for k, v in d.items()}


def to_named_tuple(d: dict, name):
    return NamedTuple(name, [(k, type(v)) for k, v in d.items()])(**d)


def create_report_doc(figs: dict = None, fname: str = None):
    from docx import Document

    figs = sorted(Path("./results").glob("*.png"))
    figs = figs[-10:]
    figs = {figs.stem: str(figs) for figs in figs}
    fname = "report"

    doc_name = f"{fname}.docx"

    document = Document()

    for title, fig in figs.items():
        title = title.replace("_", " ").title()
        document.add_heading(title, level=2)
        document.add_picture(fig)

    document.save(doc_name)

    # document.paragraphs[0].text


def custom_plot(
    sns_plots: List[Callable],
    plots_kwargs: List[Dict],
    fig_kwargs: Dict = None,
    aes_kwargs: List[Dict] = None,
    save_path: Union[str, Path] = None,
    ext: str = ".png",
    show: bool = True,
    to_pickle: bool = False,
):
    # ! TEMP
    # * just to show that it is expecting
    sns.color_palette()
    # ! TEMP

    for plot in sns_plots:
        assert (
            "ax" in inspect.signature(plot).parameters
        ), "Make sure sns_plots contain seaborn plots only"

    if fig_kwargs is None:
        fig_kwargs = {"figsize": (12, 8), "dpi": 300}

    if not sns_plots:
        raise ValueError("No plot types provided")
    elif len(sns_plots) > 1:
        # fig, ax = plt.subplots(**fig_kwargs)
        # for plot, kwargs in zip(sns_plots, plots_kwargs):
        #     plot(**kwargs, ax=ax)
        # TODO: implement this
        raise NotImplementedError("Only one plot type is supported at the moment")
    else:
        assert len(sns_plots) == len(plots_kwargs)
        if aes_kwargs is not None:
            assert isinstance(aes_kwargs, list)
            assert len(sns_plots) == len(aes_kwargs)
            aes_kwargs = aes_kwargs[0]

        sns_plot = sns_plots[0]
        plot_kwargs = plots_kwargs[0]

        fig, ax = plt.subplots(**fig_kwargs)

        sns_plot(**plot_kwargs, ax=ax)

        if aes_kwargs:
            for att, args in aes_kwargs.items():
                if hasattr(ax, att):
                    if isinstance(args, dict):
                        getattr(ax, att)(**args)
                    elif isinstance(args, list):
                        getattr(ax, att)(*args)
                    else:
                        getattr(ax, att)(args)

        # if ax.get_legend() is not None and aes_kwargs.get("legend") is None:
        #     ax.legend(bbox_to_anchor=(1, 1), loc="upper left")

    if save_path:
        image_path = Path(save_path).with_suffix(ext)
        plt.savefig(
            image_path, dpi=fig_kwargs.get("dpi", 100)
        )  # , bbox_inches="tight")
        # TODO: Check if bbox_inches="tight" does not conflict with anything else
    if to_pickle:
        if save_path is None:
            raise ValueError("if to_pickle=True, you must provide a `save_path`")
        else:
            pickle_save(save_path, fig)

    plt.show() if show else None
    plt.close()

    return fig
